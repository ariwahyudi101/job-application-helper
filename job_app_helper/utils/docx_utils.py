from __future__ import annotations

import re
from dataclasses import dataclass
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.opc.constants import RELATIONSHIP_TYPE
from docx.shared import Inches, Pt, RGBColor
from docx.text.paragraph import Paragraph

ACCENT_COLOR = RGBColor(31, 78, 95)
BODY_COLOR = RGBColor(44, 50, 56)
MUTED_COLOR = RGBColor(107, 114, 128)
PRIMARY_FONT = "Lora"
SECTION_TITLES = {
    "professional summary",
    "summary",
    "technical skills",
    "core competencies",
    "professional experience",
    "experience",
    "education",
    "professional certificate",
    "professional certificates",
    "certifications",
    "training",
    "site placement willingness",
}

INLINE_TOKEN_RE = re.compile(
    r"(\[([^\]]+)\]\(([^)]+)\)|\b(?:https?://[^\s|]+|mailto:[^\s|]+)\b|"
    r"\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Za-z]{2,}\b|"
    r"\*\*([^*]+)\*\*|_([^_]+)_|\*([^*]+)\*)"
)
HEADING_RE = re.compile(r"^(#{1,6})\s+(.*)$")


@dataclass
class ResumeHeader:
    name: str | None = None
    title: str | None = None
    contacts: list[str] | None = None


def markdown_to_docx(md_text: str, output_path: str) -> str:
    """Render resume markdown into a styled, ATS-safe DOCX."""
    doc = Document()
    _configure_page(doc)
    _configure_base_styles(doc)

    lines = [line.rstrip() for line in md_text.splitlines()]
    header, start_index = _extract_header(lines)

    if header.name or header.title or header.contacts:
        _render_header(doc, header)

    current_section: str | None = None
    current_role: str | None = None

    for raw_line in lines[start_index:]:
        line = raw_line.strip()
        if not line:
            current_role = None
            continue

        heading = _parse_heading(line)
        if heading:
            level, heading_text = heading
            if level <= 2:
                _add_section_heading(doc, heading_text)
                current_section = heading_text.casefold()
                current_role = None
                continue
            if current_section == "professional experience":
                current_role = heading_text
                _add_role_title(doc, heading_text)
                continue
            _add_subheading(doc, heading_text)
            continue

        if _is_bullet(line):
            _add_bullet(doc, line[2:].strip())
            continue

        if current_section == "professional experience":
            if _is_role_line(line):
                current_role = _strip_wrapping_emphasis(line)
                _add_role_title(doc, current_role)
                continue
            if current_role and _looks_like_metadata(line):
                _add_metadata_line(doc, line)
                continue

        if _looks_like_metadata(line) and current_role:
            _add_metadata_line(doc, line)
            continue

        _add_body_paragraph(doc, line)

    output = Path(output_path)
    output.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output))
    return str(output)


def _configure_page(doc: Document) -> None:
    section = doc.sections[0]
    section.start_type = WD_SECTION.CONTINUOUS
    section.top_margin = Inches(0.55)
    section.bottom_margin = Inches(0.55)
    section.left_margin = Inches(0.7)
    section.right_margin = Inches(0.7)


def _configure_base_styles(doc: Document) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = PRIMARY_FONT
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = BODY_COLOR


def _extract_header(lines: list[str]) -> tuple[ResumeHeader, int]:
    header = ResumeHeader(contacts=[])
    non_empty = [(index, line.strip()) for index, line in enumerate(lines) if line.strip()]
    if not non_empty:
        return header, 0

    cursor = 0

    _, first_line = non_empty[cursor]
    first_heading = _parse_heading(first_line)
    if first_heading and first_heading[0] == 1:
        header.name = first_heading[1]
        cursor += 1
    elif _looks_like_name_line(first_line):
        header.name = _strip_wrapping_emphasis(first_line)
        cursor += 1

    if cursor < len(non_empty):
        _, second_line = non_empty[cursor]
        second_heading = _parse_heading(second_line)
        if second_heading and second_heading[0] == 2:
            header.title = second_heading[1]
            cursor += 1
        elif not _parse_section_title(second_line) and not _is_bullet(second_line):
            if _looks_like_title_line(second_line):
                header.title = _strip_wrapping_emphasis(second_line)
                cursor += 1

    if cursor < len(non_empty):
        _, third_line = non_empty[cursor]
        third_heading = _parse_heading(third_line)
        contact_candidate = third_heading[1] if third_heading and third_heading[0] >= 3 else third_line
        if _looks_like_contact_line(contact_candidate):
            header.contacts = [
                piece.strip() for piece in contact_candidate.split("|") if piece.strip()
            ]
            cursor += 1

    start_index = non_empty[cursor][0] if cursor < len(non_empty) else len(lines)
    return header, start_index


def _render_header(doc: Document, header: ResumeHeader) -> None:
    if header.name:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(2)
        run = p.add_run(header.name)
        run.font.name = PRIMARY_FONT
        run.font.size = Pt(30)
        run.font.bold = True
        run.font.color.rgb = ACCENT_COLOR

    if header.title:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(8)
        run = p.add_run(header.title)
        run.font.name = PRIMARY_FONT
        run.font.size = Pt(13.5)
        run.font.bold = True
        run.font.color.rgb = RGBColor(66, 84, 95)

    if header.contacts:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(14)
        p.paragraph_format.line_spacing = 1.0
        for index, item in enumerate(header.contacts):
            if index:
                sep = p.add_run("  |  ")
                sep.font.size = Pt(10)
                sep.font.name = PRIMARY_FONT
                sep.font.color.rgb = MUTED_COLOR
            _append_inline_runs(p, item, contact_mode=True)

    divider = doc.add_paragraph()
    divider.paragraph_format.space_after = Pt(16)
    _set_bottom_border(divider, "D8DEE6", 8)


def _parse_section_title(line: str) -> str | None:
    heading = _parse_heading(line)
    if heading and heading[0] <= 2:
        return heading[1]

    clean = _strip_wrapping_emphasis(line)
    if clean.casefold() in SECTION_TITLES:
        return clean.title()
    if clean.isupper() and len(clean.split()) <= 5:
        return clean.title()
    return None


def _parse_heading(line: str) -> tuple[int, str] | None:
    match = HEADING_RE.match(line.strip())
    if not match:
        return None
    return len(match.group(1)), match.group(2).strip()


def _add_section_heading(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(7)
    run = p.add_run(text.upper())
    run.font.name = PRIMARY_FONT
    run.font.size = Pt(11.5)
    run.font.bold = True
    run.font.color.rgb = ACCENT_COLOR
    _set_bottom_border(p, "C7D4DC", 10)


def _add_role_title(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(text)
    run.font.name = PRIMARY_FONT
    run.font.size = Pt(11.5)
    run.font.bold = True
    run.font.color.rgb = BODY_COLOR


def _add_subheading(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(text)
    run.font.name = PRIMARY_FONT
    run.font.size = Pt(10.8)
    run.font.bold = True
    run.font.color.rgb = BODY_COLOR


def _add_metadata_line(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing = 1.0
    _append_inline_runs(p, text, metadata_mode=True)


def _add_body_paragraph(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.08
    _append_inline_runs(p, text)


def _add_bullet(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.left_indent = Inches(0.2)
    p.paragraph_format.first_line_indent = Inches(-0.15)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing = 1.05

    bullet = p.add_run("• ")
    bullet.font.name = PRIMARY_FONT
    bullet.font.size = Pt(10.5)
    bullet.font.color.rgb = ACCENT_COLOR

    _append_inline_runs(p, text)


def _append_inline_runs(
    paragraph: Paragraph,
    text: str,
    *,
    contact_mode: bool = False,
    metadata_mode: bool = False,
) -> None:
    pos = 0
    for match in INLINE_TOKEN_RE.finditer(text):
        raw_token = match.group(0)
        if match.start() > pos:
            _add_text_run(
                paragraph,
                text[pos:match.start()],
                contact_mode=contact_mode,
                metadata_mode=metadata_mode,
            )

        if match.group(2) is not None and match.group(3) is not None:
            url = match.group(3).strip()
            label = match.group(2).strip()
            _add_hyperlink(paragraph, url, label, contact_mode=contact_mode)
        elif raw_token.startswith(("http://", "https://", "mailto:")) or "@" in raw_token:
            raw_link = raw_token.strip()
            if raw_link.startswith(("http://", "https://", "mailto:")):
                label = raw_link.replace("mailto:", "")
                _add_hyperlink(paragraph, raw_link, label, contact_mode=contact_mode)
            elif "@" in raw_link:
                _add_hyperlink(paragraph, f"mailto:{raw_link}", raw_link, contact_mode=contact_mode)
        elif match.group(4) is not None:
            _add_text_run(
                paragraph,
                match.group(4),
                bold=True,
                contact_mode=contact_mode,
                metadata_mode=metadata_mode,
            )
        else:
            italic_text = match.group(5) if match.group(5) is not None else match.group(6)
            _add_text_run(
                paragraph,
                italic_text,
                italic=True,
                contact_mode=contact_mode,
                metadata_mode=metadata_mode,
            )
        pos = match.end()

    if pos < len(text):
        _add_text_run(
            paragraph,
            text[pos:],
            contact_mode=contact_mode,
            metadata_mode=metadata_mode,
        )


def _add_text_run(
    paragraph: Paragraph,
    text: str,
    *,
    bold: bool = False,
    italic: bool = False,
    contact_mode: bool = False,
    metadata_mode: bool = False,
) -> None:
    if not text:
        return

    run = paragraph.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.name = PRIMARY_FONT
    run.font.size = Pt(10 if contact_mode else 10 if metadata_mode else 10.5)
    run.font.color.rgb = MUTED_COLOR if metadata_mode else BODY_COLOR


def _add_hyperlink(
    paragraph: Paragraph,
    url: str,
    label: str,
    *,
    contact_mode: bool = False,
) -> None:
    part = paragraph.part
    r_id = part.relate_to(url, RELATIONSHIP_TYPE.HYPERLINK, is_external=True)

    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)

    new_run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")

    color = OxmlElement("w:color")
    color.set(qn("w:val"), "19558E")
    r_pr.append(color)

    fonts = OxmlElement("w:rFonts")
    fonts.set(qn("w:ascii"), PRIMARY_FONT)
    fonts.set(qn("w:hAnsi"), PRIMARY_FONT)
    fonts.set(qn("w:cs"), PRIMARY_FONT)
    r_pr.append(fonts)

    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    r_pr.append(underline)

    font_size = OxmlElement("w:sz")
    font_size.set(qn("w:val"), "19" if contact_mode else "21")
    r_pr.append(font_size)

    new_run.append(r_pr)
    text = OxmlElement("w:t")
    text.text = label
    new_run.append(text)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


def _set_bottom_border(paragraph: Paragraph, color: str, size: int) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    borders = p_pr.find(qn("w:pBdr"))
    if borders is None:
        borders = OxmlElement("w:pBdr")
        p_pr.append(borders)

    bottom = borders.find(qn("w:bottom"))
    if bottom is None:
        bottom = OxmlElement("w:bottom")
        borders.append(bottom)

    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(size))
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), color)


def _looks_like_name_line(line: str) -> bool:
    clean = _strip_wrapping_emphasis(line)
    if not clean:
        return False
    if any(token in clean for token in ("|", "@", "http://", "https://", "mailto:")):
        return False
    words = clean.split()
    return 1 < len(words) <= 5 and not clean.isupper()


def _looks_like_contact_line(line: str) -> bool:
    clean = _strip_wrapping_emphasis(line)
    if not clean:
        return False
    parts = [part.strip() for part in clean.split("|") if part.strip()]
    if not parts:
        return False
    contact_like_parts = sum(1 for part in parts if _is_contact_fragment(part))
    if len(parts) == 1:
        return contact_like_parts == 1
    return contact_like_parts >= 2


def _looks_like_title_line(line: str) -> bool:
    clean = _strip_wrapping_emphasis(line)
    if not clean or _parse_section_title(clean) or _is_bullet(clean):
        return False
    if _looks_like_contact_line(clean):
        return False
    return len(clean.split()) <= 12


def _is_bullet(line: str) -> bool:
    return line.startswith("- ") or line.startswith("* ")


def _is_role_line(line: str) -> bool:
    clean = _strip_wrapping_emphasis(line)
    if not clean or ":" in clean or "|" in clean:
        return False
    return line.startswith("**") and line.endswith("**")


def _looks_like_metadata(line: str) -> bool:
    clean = _strip_wrapping_emphasis(line)
    return "|" in clean or clean.startswith("Industry:")


def _strip_wrapping_emphasis(text: str) -> str:
    stripped = text.strip()
    if stripped.startswith("**") and stripped.endswith("**") and len(stripped) > 4:
        return stripped[2:-2].strip()
    if stripped.startswith("_") and stripped.endswith("_") and len(stripped) > 2:
        return stripped[1:-1].strip()
    return stripped


def _is_contact_fragment(text: str) -> bool:
    clean = text.strip()
    if not clean:
        return False
    if re.search(r"https?://|mailto:|@", clean):
        return True
    digit_count = sum(1 for char in clean if char.isdigit())
    return digit_count >= 6
